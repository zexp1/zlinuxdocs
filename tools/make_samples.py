#!/usr/bin/env python3
"""Generate the example documents that ship with zlinuxdocs.

These are written from scratch with python-docx. Nothing here is copied from
anybody's real work product.

Usage:  python3 tools/make_samples.py [outdir]
"""

import os
import sys

HERE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(HERE, "vendor"))
sys.path.insert(0, HERE)

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt  # noqa: E402

from zlinuxdocs.lib.docxio import ensure_style  # noqa: E402

BODY = (
    "This paragraph exists so the document has enough real text for a PDF "
    "text-layer check to mean something. It describes an imaginary district "
    "programme, its budget lines, its timelines and its expected outcomes, in "
    "the flat register that official reports are written in. Nothing in it "
    "refers to any real place, person or scheme."
)

BODY2 = (
    "The costing below is illustrative only. Figures are rounded to the nearest "
    "lakh and carry no commitment. Where a component depends on a procurement "
    "that has not yet been tendered, the estimate is marked provisional and is "
    "expected to move within a band of ten per cent either way."
)


def _pseudo(doc, text, bold=True, size=None):
    """A paragraph that LOOKS like a heading but is Normal style."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p


def plain_no_headings(path):
    doc = Document()
    _pseudo(doc, "District Connectivity Programme", size=20)
    _pseudo(doc, "Detailed Project Report - Volume I", bold=False, size=13)
    doc.add_paragraph("Prepared for internal review. Not for circulation.")

    _pseudo(doc, "1 - Introduction", size=16)
    doc.add_paragraph(BODY)
    _pseudo(doc, "1.1 Scope", size=14)
    doc.add_paragraph(BODY2)
    _pseudo(doc, "1.2 Method", size=14)
    doc.add_paragraph(BODY)
    _pseudo(doc, "1.2.1 Field survey", size=12)
    doc.add_paragraph(BODY2)
    _pseudo(doc, "1.2.2 Desk review", size=12)
    doc.add_paragraph(BODY)

    _pseudo(doc, "2 - Present position", size=16)
    doc.add_paragraph(BODY2)
    _pseudo(doc, "2.1 Coverage today", size=14)
    doc.add_paragraph(BODY)
    _pseudo(doc, "2.2 Gaps", size=14)
    doc.add_paragraph(BODY2)

    _pseudo(doc, "3 - Proposal", size=16)
    doc.add_paragraph(BODY)
    _pseudo(doc, "3.1 Components", size=14)
    doc.add_paragraph(BODY2)
    _pseudo(doc, "3.2 Phasing", size=14)
    doc.add_paragraph(BODY)

    _pseudo(doc, "ANNEXURE I", size=16)
    doc.add_paragraph(BODY2)
    doc.add_paragraph("See 1.1 above for the scope note; this line must NOT be promoted.")
    doc.add_paragraph(
        "A long sentence that begins with numbers 1.2 but then runs on far past any "
        "reasonable heading length, describing at tedious length the way in which "
        "estimates were reconciled against the previous financial year, so that the "
        "detector has something it must refuse to promote."
    )
    doc.save(path)


def proper_headings(path):
    doc = Document()
    ensure_style(doc, "Heading 1", outline_level=0)
    ensure_style(doc, "Heading 2", outline_level=1)
    ensure_style(doc, "Heading 3", outline_level=2)

    doc.add_paragraph("District Connectivity Programme", style="Title")
    doc.add_paragraph("Detailed Project Report - Volume I", style="Subtitle")

    doc.add_paragraph("1 - Introduction", style="Heading 1")
    doc.add_paragraph(BODY)
    doc.add_paragraph("1.1 Scope", style="Heading 2")
    doc.add_paragraph(BODY2)
    doc.add_paragraph("1.1.1 Boundaries", style="Heading 3")
    doc.add_paragraph(BODY)
    doc.add_paragraph("2 - Present position", style="Heading 1")
    doc.add_paragraph(BODY2)
    doc.add_paragraph("2.1 Coverage today", style="Heading 2")
    doc.add_paragraph(BODY)
    doc.add_paragraph("3 - Proposal", style="Heading 1")
    doc.add_paragraph(BODY2)
    doc.add_paragraph("3.1 Components", style="Heading 2")
    doc.add_paragraph(BODY)
    doc.add_paragraph("3.2 Phasing", style="Heading 2")
    doc.add_paragraph(BODY2)
    doc.save(path)


def mixed_tables(path):
    doc = Document()
    ensure_style(doc, "Heading 1", outline_level=0)
    ensure_style(doc, "Heading 2", outline_level=1)

    doc.add_paragraph("Costing Statement", style="Title")
    doc.add_paragraph("1 - Summary of costs", style="Heading 1")
    doc.add_paragraph(BODY)

    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, label in enumerate(["Sl.", "Component", "Unit cost (Rs lakh)", "Total (Rs lakh)"]):
        hdr[i].text = label
    for i in range(1, 41):
        row = t.add_row().cells
        row[0].text = str(i)
        row[1].text = "Component %02d - civil works and commissioning" % i
        row[2].text = "%0.2f" % (12.5 + i)
        row[3].text = "%0.2f" % ((12.5 + i) * 3)

    doc.add_paragraph("2 - Notes on the table", style="Heading 1")
    doc.add_paragraph(BODY2)
    doc.add_paragraph("2.1 Assumptions", style="Heading 2")
    doc.add_paragraph(BODY)

    t2 = doc.add_table(rows=2, cols=2)
    t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Assumption"
    t2.rows[0].cells[1].text = "Basis"
    t2.rows[1].cells[0].text = "Ten per cent contingency"
    t2.rows[1].cells[1].text = "Standard departmental practice"
    doc.save(path)


def with_markers(path):
    doc = Document()
    ensure_style(doc, "Heading 1", outline_level=0)
    ensure_style(doc, "Heading 2", outline_level=1)

    doc.add_paragraph("Draft Under Review", style="Title")
    doc.add_paragraph("1 - Background", style="Heading 1")
    doc.add_paragraph(BODY)
    doc.add_paragraph("The 2023 baseline figure of 41 per cent is quoted from memory. [V]")
    doc.add_paragraph("1.1 Population served", style="Heading 2")
    doc.add_paragraph(BODY2)
    doc.add_paragraph("Check whether the block boundary changed after the 2024 notification. [V]")

    doc.add_paragraph("2 - Costing", style="Heading 1")
    t = doc.add_table(rows=3, cols=3)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Item"
    t.rows[0].cells[1].text = "Amount"
    t.rows[0].cells[2].text = "Remarks"
    t.rows[1].cells[0].text = "Civil works"
    t.rows[1].cells[1].text = "126.00"
    t.rows[1].cells[2].text = "Rate to be re-confirmed with PWD [V]"
    t.rows[2].cells[0].text = "Equipment"
    t.rows[2].cells[1].text = "84.50"
    t.rows[2].cells[2].text = "Quotation is over a year old [V]"

    doc.add_paragraph("3 - Recommendation", style="Heading 1")
    doc.add_paragraph(BODY)
    doc.add_paragraph("Signature block to be confirmed with the Director's office. [V]")
    doc.save(path)


def unicode_heavy(path):
    doc = Document()
    ensure_style(doc, "Heading 1", outline_level=0)
    ensure_style(doc, "Heading 2", outline_level=1)

    doc.add_paragraph("বহুভাষিক নথি পরীক্ষা", style="Title")
    doc.add_paragraph("Multilingual document test", style="Subtitle")

    doc.add_paragraph("1 - বাংলা অংশ", style="Heading 1")
    doc.add_paragraph(
        "এই অনুচ্ছেদটি বাংলা লিপিতে লেখা হয়েছে যাতে পরীক্ষা করা যায় যে নথি "
        "রূপান্তরের পরে অক্ষরগুলি অবিকৃত থাকে। কোনো তথ্য বাস্তব নয়।"
    )
    doc.add_paragraph("1.1 তালিকা", style="Heading 2")
    doc.add_paragraph("প্রথম · দ্বিতীয় · তৃতীয় · চতুর্থ · পঞ্চম")

    doc.add_paragraph("2 - Kokborok kokrok", style="Heading 1")
    doc.add_paragraph(
        "Kokborok bo Tripura ni official kok. Bono romani hamjak kaham thangnai "
        "kolopni bagwi kaisa lekha tongo. Bosorok chini bufang hamjaklai."
    )
    doc.add_paragraph("2.1 Nokhtai", style="Heading 2")
    doc.add_paragraph("Sa · Nwi · Tham · Brwi · Ba · Dok · Sni")

    doc.add_paragraph("3 - देवनागरी खंड", style="Heading 1")
    doc.add_paragraph(
        "यह अनुच्छेद देवनागरी लिपि में लिखा गया है ताकि यह जाँचा जा सके कि "
        "रूपांतरण के बाद अक्षर सही बने रहते हैं। यहाँ कोई वास्तविक आँकड़ा नहीं है।"
    )
    doc.add_paragraph("4 - Symbols and dashes", style="Heading 1")
    doc.add_paragraph("Rs 1,20,000 — “curly quotes” · ½ ¾ ° ± × ÷ … ‰ € ₹ £ ¥")
    doc.save(path)


def empty_doc(path):
    doc = Document()
    doc.save(path)


def one_line(path):
    doc = Document()
    p = doc.add_paragraph("Received. Please acknowledge.")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.save(path)


def weird_name(path):
    doc = Document()
    _pseudo(doc, "Final Submission (Revised)", size=20)
    _pseudo(doc, "1 - Covering note", size=16)
    doc.add_paragraph(BODY)
    _pseudo(doc, "1.1 Enclosures", size=14)
    doc.add_paragraph(BODY2)
    doc.add_paragraph("One item still needs a check before dispatch. [V]")
    doc.save(path)


SAMPLES = [
    ("plain-no-headings.docx", plain_no_headings),
    ("proper-headings.docx", proper_headings),
    ("mixed-tables.docx", mixed_tables),
    ("with-markers.docx", with_markers),
    ("unicode-heavy.docx", unicode_heavy),
    ("empty.docx", empty_doc),
    ("one-line.docx", one_line),
    ("weird name (final) v2.docx", weird_name),
]


def main():
    outdir = sys.argv[1] if len(sys.argv) > 1 else os.path.join(HERE, "share", "samples")
    os.makedirs(outdir, exist_ok=True)
    for name, fn in SAMPLES:
        path = os.path.join(outdir, name)
        fn(path)
        print("wrote %s (%d bytes)" % (path, os.path.getsize(path)))
    return 0


if __name__ == "__main__":
    sys.exit(main())
